#!/usr/bin/env python3
"""
inspect_format.py — deterministic ATS format inspector for .docx / .pdf resumes.

Part of the ats-resume-checker skill. Detects structural elements that commonly
break Applicant Tracking System (ATS) parsing: multi-column layouts, tables,
text boxes, images/graphics, non-standard fonts, and contact info stranded in
headers/footers. Reading extracted *text* alone can't see these — this script
inspects the actual document structure.

Usage:
    python inspect_format.py <path-to-resume.(docx|pdf)>          # inspect layout
    python inspect_format.py --text <path-to-resume.(docx|pdf)>   # extract body text

The two modes fail in opposite ways, and callers branch on the difference:

    inspect  -> always exits 0, and reports trouble as {"available": false, ...}
    --text   -> writes the text to stdout, or the reason to stderr and exits 1

Keep that contract. The skill's step 2 recovery path keys on `--text` exiting
non-zero, and would silently stop firing if this mode started exiting 0.

Output (inspect mode): a single JSON object on stdout:
    {
      "available": true|false,   # false => deps missing / unsupported type => caller falls back to prompt-only heuristics
      "reason": "...",           # present only when available is false
      "file_type": "docx"|"pdf",
      "facts": {...},            # raw detected structure
      "flags": [ {"level": "fail"|"warn"|"pass", "category": "layout", "message": "..."} ]
    }

Design notes:
  - Never raises on a malformed file — always prints one JSON object.
  - Dependencies (python-docx, pdfplumber) are optional and imported lazily;
    if missing, "available": false tells the skill to use prompt-only checks.
"""
import json
import os
import sys

# Web-safe fonts widely parsed by ATS (see skill research notes).
SAFE_FONTS = {
    "Arial", "Calibri", "Cambria", "Garamond", "Georgia",
    "Helvetica", "Tahoma", "Times New Roman", "Verdana",
}
_SAFE_NORM = {f.replace(" ", "").lower() for f in SAFE_FONTS}

# Bullet glyphs ATS parsers handle reliably. Anything else leading a line
# (✓ ➢ ★ ) is a custom symbol that commonly drops out or turns into mojibake.
STANDARD_BULLETS = set("•◦▪§·-–—*")


def _base_font(name):
    """Normalize a raw font name for comparison.

    PDF font names arrive subset-prefixed and style-suffixed
    ('BAAAAA+Carlito-Bold', 'TimesNewRomanPSMT'); docx names are usually clean.
    """
    if "+" in name:
        name = name.split("+", 1)[1]
    name = name.split("-", 1)[0].split(",", 1)[0]
    return name.strip()


def _is_safe_font(name):
    n = _base_font(name).replace(" ", "").lower()
    return any(n.startswith(s) for s in _SAFE_NORM)


def _unsafe_fonts(fonts):
    """Distinct display names of fonts outside SAFE_FONTS, deduped by base name."""
    return sorted({_base_font(f) for f in fonts if not _is_safe_font(f)})


def _split_decoration_fonts(font_chars):
    """Separate fonts that carry words from fonts that only draw marks.

    (content_fonts, decoration_display_names)

    A font rendering no letters or digits is decoration — the bullet glyph, a
    rule, a tick. Word exports the round bullet in Symbol, so nearly every
    resume PDF embeds one, and judging it as a document font reports a
    non-web-safe font on a document whose every word is Calibri. A font
    carrying actual text still reports, because its characters are alnum.
    """
    content = {f for f, chars in font_chars.items() if any(c.isalnum() for c in chars)}
    decoration = sorted({_base_font(f) for f in font_chars if f not in content})
    return content, decoration


def _bullet_flag(lines):
    """Flag non-standard bullet glyphs leading any line.

    Returns None when the document has no bullets at all — there is nothing to
    pass, and a spurious pass becomes a false ✅ in the report.
    """
    odd, seen = set(), False
    for line in lines:
        s = line.strip()
        if not s:
            continue
        ch = s[0]
        if ch.isalnum() or ch in "\"'([{#/&":
            continue
        if ch in STANDARD_BULLETS:
            seen = True
        else:
            odd.add(ch)
    if odd:
        return {"level": "warn", "category": "layout",
                "message": f"Non-standard bullet glyph(s) {' '.join(sorted(odd))} — "
                           "custom symbols are often dropped or mangled by ATS; "
                           "use a plain round bullet or a hyphen."}
    if seen:
        return {"level": "pass", "category": "layout", "message": "Standard bullet characters."}
    return None


def _docx_report(path):
    try:
        import docx  # python-docx
        from docx.oxml.ns import qn
    except ImportError:
        return {"available": False, "reason": "python-docx not installed (python3 -m pip install python-docx)"}

    try:
        doc = docx.Document(path)
    except Exception as e:  # noqa: BLE001 - report, never crash
        return {"available": False, "reason": f"could not open docx: {e}"}

    facts, flags = {}, []
    body_xml = doc.element.body.xml

    # --- Tables ---
    n_tables = len(doc.tables)
    facts["tables"] = n_tables
    if n_tables:
        flags.append({"level": "fail", "category": "layout",
                      "message": f"{n_tables} table(s) found — a real structural element in .docx, and ATS "
                                 "often scrambles or drops content inside tables."})
    else:
        flags.append({"level": "pass", "category": "layout", "message": "No tables."})

    # --- Multi-column layout (sectPr/w:cols[@w:num]) ---
    max_cols = 1
    for section in doc.sections:
        cols = section._sectPr.find(qn("w:cols"))
        if cols is not None:
            num = cols.get(qn("w:num"))
            if num:
                try:
                    max_cols = max(max_cols, int(num))
                except ValueError:
                    pass
    facts["columns"] = max_cols
    if max_cols > 1:
        flags.append({"level": "fail", "category": "layout",
                      "message": f"{max_cols}-column layout — multi-column resumes are frequently read out of order by ATS."})
    else:
        flags.append({"level": "pass", "category": "layout", "message": "Single-column layout."})

    # --- Text boxes ---
    has_textbox = ("txbxContent" in body_xml) or ("v:textbox" in body_xml)
    facts["text_boxes"] = has_textbox
    if has_textbox:
        flags.append({"level": "fail", "category": "layout",
                      "message": "Text box(es) detected — text inside boxes is commonly ignored by ATS."})
    else:
        flags.append({"level": "pass", "category": "layout", "message": "No text boxes."})

    # --- Images / graphics ---
    n_inline = len(doc.inline_shapes)
    has_graphics = ("<w:drawing" in body_xml) or ("graphicData" in body_xml) or ("<pic:pic" in body_xml)
    facts["inline_images"] = n_inline
    facts["has_graphics"] = has_graphics
    if n_inline or has_graphics:
        flags.append({"level": "warn", "category": "layout",
                      "message": "Images/graphics detected — any text baked into them is invisible to ATS."})
    else:
        flags.append({"level": "pass", "category": "layout", "message": "No images or graphics."})

    # --- Header/footer text (contact-info risk) ---
    hf_text = []
    for section in doc.sections:
        for part in (section.header, section.footer):
            try:
                t = "\n".join(p.text for p in part.paragraphs).strip()
                if t:
                    hf_text.append(t)
            except Exception:  # noqa: BLE001
                pass
    facts["header_footer_text"] = hf_text
    if hf_text:
        flags.append({"level": "warn", "category": "layout",
                      "message": "Text found in header/footer — in a .docx this sits outside the main document body, and whether a parser reads it depends on the extractor; keep contact info in the body."})
    else:
        flags.append({"level": "pass", "category": "layout",
                      "message": "No header/footer text — contact details are in the body."})

    # --- Fonts ---
    # Run-level names only exist where the font was set explicitly; runs that
    # inherit from a style report None, so also read the document defaults.
    fonts = set()
    for p in doc.paragraphs:
        for r in p.runs:
            if r.font is not None and r.font.name:
                fonts.add(r.font.name)
    for style_name in ("Normal", "Body Text"):
        try:
            st = doc.styles[style_name]
            if st.font is not None and st.font.name:
                fonts.add(st.font.name)
        except KeyError:
            pass
    facts["fonts"] = sorted(fonts)
    unsafe = _unsafe_fonts(fonts)
    if unsafe:
        flags.append({"level": "warn", "category": "layout",
                      "message": f"Non-standard font(s): {', '.join(unsafe)} — prefer web-safe fonts (Arial, Calibri, Times New Roman, ...)."})
    elif fonts:
        flags.append({"level": "pass", "category": "layout", "message": "Web-safe fonts only."})
    else:
        # Empty is not clean — it means every run inherited and nothing was readable.
        flags.append({"level": "warn", "category": "layout",
                      "message": "Font information unavailable — fonts are set at the theme level, "
                                 "so this check could not run. Verify manually that the font is web-safe."})

    # --- Bullets and length ---
    para_text = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                para_text.extend(p.text for p in cell.paragraphs)
    bf = _bullet_flag(para_text)
    if bf:
        flags.append(bf)
    facts["word_count"] = len(" ".join(para_text).split())
    flags.append({"level": "pass", "category": "layout",
                  "message": ".docx — a natively parseable format."})

    return {"available": True, "file_type": "docx", "facts": facts, "flags": flags}


def _pdf_report(path):
    try:
        import pdfplumber
    except ImportError:
        return {"available": False, "reason": "pdfplumber not installed (python3 -m pip install pdfplumber)"}

    try:
        pdf = pdfplumber.open(path)
    except Exception as e:  # noqa: BLE001
        return {"available": False, "reason": f"could not open pdf: {e}"}

    facts, flags = {}, []
    n_pages = n_images = n_tables = multicol_pages = 0
    font_chars = {}
    text_parts = []

    try:
        for page in pdf.pages:
            n_pages += 1
            n_images += len(page.images or [])
            text_parts.append(page.extract_text() or "")
            try:
                n_tables += len(page.find_tables() or [])
            except Exception:  # noqa: BLE001
                pass
            for ch in (page.chars or []):
                fn = ch.get("fontname")
                if fn:
                    font_chars.setdefault(fn, set()).add(ch.get("text", ""))
            # Crude multi-column heuristic: central gutter empty while both sides carry text.
            try:
                w = page.width or 0
                left = right = center = 0
                for ch in (page.chars or []):
                    xm = (ch["x0"] + ch["x1"]) / 2.0
                    frac = (xm / w) if w else 0
                    if frac < 0.45:
                        left += 1
                    elif frac > 0.55:
                        right += 1
                    else:
                        center += 1
                total = left + right + center
                if total > 200 and left > 0.2 * total and right > 0.2 * total and center < 0.05 * total:
                    multicol_pages += 1
            except Exception:  # noqa: BLE001
                pass
    finally:
        pdf.close()

    text = "\n".join(text_parts)
    facts.update({
        "pages": n_pages,
        "images": n_images,
        "tables_detected": n_tables,
        "fonts": sorted(font_chars),
        "multi_column_pages": multicol_pages,
        "extracted_chars": len(text.strip()),
        "word_count": len(text.split()),
    })

    # --- Scanned / image-only PDF: no extractable text at all ---
    # The most fatal ATS failure there is, and invisible from extracted text
    # (there isn't any). Checked first because nothing else below is meaningful.
    if facts["extracted_chars"] < 50:
        flags.append({"level": "fail", "category": "layout",
                      "message": "No extractable text — this is a scanned or image-only PDF. "
                                 "ATS cannot read a single word of it. Export a text-based PDF "
                                 "from the original document, or submit .docx."})
        return {"available": True, "file_type": "pdf", "facts": facts, "flags": flags}

    if multicol_pages:
        flags.append({"level": "fail", "category": "layout",
                      "message": f"Multi-column layout on {multicol_pages} page(s) — likely read out of order by ATS."})
    else:
        flags.append({"level": "pass", "category": "layout", "message": "No multi-column layout detected."})

    # warn, where the docx path fails: the asymmetry is deliberate. A PDF holds no
    # table objects at all, so find_tables() infers one from ruling lines and
    # false-positives on any bordered box. The level reflects confidence in the
    # detection, not a smaller risk — a real table parses just as badly here.
    if n_tables:
        flags.append({"level": "warn", "category": "layout",
                      "message": f"{n_tables} table-like structure(s) detected — inferred from ruling lines, "
                                 "so it may be a false positive. A PDF stores positioned text rather than "
                                 "table objects, so this is a detection-confidence warning, not a smaller "
                                 "risk: if content really is in a table, the parse risk matches a .docx. "
                                 "Verify no content is trapped in a table."})
    else:
        flags.append({"level": "pass", "category": "layout", "message": "No tables detected."})

    if n_images:
        flags.append({"level": "warn", "category": "layout",
                      "message": f"{n_images} image(s) — text inside images is invisible to ATS."})
    else:
        flags.append({"level": "pass", "category": "layout", "message": "No images."})

    flags.append({"level": "pass", "category": "layout",
                  "message": "Text-based PDF — text extracts cleanly, no OCR needed."})

    content_fonts, decoration = _split_decoration_fonts(font_chars)
    facts["decoration_fonts"] = decoration
    # Only worth naming the ones that would otherwise have been reported: a
    # web-safe font drawing the spaces after bullets is not news.
    noted = [f for f in decoration if not _is_safe_font(f)]
    aside = ""
    if noted:
        verb = "is" if len(noted) == 1 else "are"
        aside = (f" {', '.join(noted)} {verb} also embedded, but draws only bullet "
                 "glyphs and punctuation, never words.")
    unsafe = _unsafe_fonts(content_fonts)
    if unsafe:
        flags.append({"level": "warn", "category": "layout",
                      "message": f"Non-standard font(s) in the text: {', '.join(unsafe)} — prefer web-safe fonts (Arial, Calibri, Times New Roman, ...).{aside}"})
    elif content_fonts:
        flags.append({"level": "pass", "category": "layout",
                      "message": f"Web-safe fonts throughout the text.{aside}"})
    else:
        flags.append({"level": "warn", "category": "layout",
                      "message": "Font information unavailable — the PDF exposes no readable font "
                                 "names, so this check could not run. Verify manually."})

    bf = _bullet_flag(text.split("\n"))
    if bf:
        flags.append(bf)

    return {"available": True, "file_type": "pdf", "facts": facts, "flags": flags}


def _docx_text(path):
    """Body text in reading order, including table cells. (text, error)

    `Document.paragraphs` skips everything inside a table, which is exactly
    where a badly built resume puts its content — so extracting from it alone
    reports sections as missing when they are on the page. Walking the body
    element keeps paragraphs and table cells in the order they appear.

    Header and footer text is deliberately excluded. The inspector reports it
    separately in `facts.header_footer_text`, and Category 5 grades a contact
    detail found there as misplaced rather than absent; folding it into the body
    text would erase the distinction the grader needs to tell those apart.
    """
    try:
        import docx  # python-docx
        from docx.oxml.ns import qn
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except ImportError:
        return None, "python-docx not installed (python3 -m pip install python-docx)"

    try:
        doc = docx.Document(path)
    except Exception as e:  # noqa: BLE001 - report, never crash
        return None, f"could not open docx: {e}"

    lines = []
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            text = Paragraph(child, doc).text.strip()
            if text:
                lines.append(text)
        elif child.tag == qn("w:tbl"):
            for row in Table(child, doc).rows:
                # Merged cells repeat their text across the span; keep one copy.
                seen, cells = set(), []
                for cell in row.cells:
                    text = cell.text.strip()
                    if text and text not in seen:
                        seen.add(text)
                        cells.append(text)
                if cells:
                    lines.append(" | ".join(cells))
    return "\n".join(lines), None


def _pdf_text(path):
    """Page text in reading order. (text, error)"""
    try:
        import pdfplumber
    except ImportError:
        return None, "pdfplumber not installed (python3 -m pip install pdfplumber)"
    try:
        with pdfplumber.open(path) as pdf:
            return "\n".join((page.extract_text() or "") for page in pdf.pages), None
    except Exception as e:  # noqa: BLE001 - report, never crash
        return None, f"could not open pdf: {e}"


def main():
    args = [a for a in sys.argv[1:] if a != "--text"]
    text_mode = "--text" in sys.argv[1:]
    usage = "usage: python inspect_format.py [--text] <resume.docx|resume.pdf>"

    if len(args) != 1:
        if text_mode:
            print(usage, file=sys.stderr)
            sys.exit(1)
        print(json.dumps({"available": False, "reason": usage}))
        return
    path = args[0]
    if not os.path.isfile(path):
        if text_mode:
            print(f"file not found: {path}", file=sys.stderr)
            sys.exit(1)
        print(json.dumps({"available": False, "reason": f"file not found: {path}"}))
        return
    ext = os.path.splitext(path)[1].lower()

    if text_mode:
        if ext == ".docx":
            text, err = _docx_text(path)
        elif ext == ".pdf":
            text, err = _pdf_text(path)
        else:
            print(f"--text handles .docx and .pdf only; read {ext} files directly.", file=sys.stderr)
            sys.exit(1)
        if err:
            print(err, file=sys.stderr)
            sys.exit(1)
        print(text)
        return

    if ext == ".docx":
        out = _docx_report(path)
    elif ext == ".pdf":
        out = _pdf_report(path)
    else:
        out = {"available": False,
               "reason": f"unsupported file type '{ext}' — only .docx and .pdf are inspected structurally; use prompt-only heuristics for text/markdown."}
    print(json.dumps(out, indent=2))


if __name__ == "__main__":
    main()
